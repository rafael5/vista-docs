#!/usr/bin/env python3
"""
corpus_survey.py — Empirical element-type frequency analysis across VDL DOCX corpus.

Reads every fetched .docx file referenced in the manifest (or a raw/ directory)
and produces a structured survey of:

  1. Paragraph styles          — frequency table across all docs
  2. Table header patterns     — top-N column header combinations → table type signal
  3. Pseudo-heading candidates — bold standalone lines not styled as Heading N
  4. Cross-reference patterns  — internal/external page refs, "see section" patterns
  5. Special block patterns    — NOTE:/WARNING:/CAUTION: callouts, terminal sessions
  6. Per-document summary      — section count, table count, pseudo-heading count
  7. Corpus-level gaps         — styles seen in only one doc (one-offs vs. standards)

Output:
  survey-report.txt    — human-readable report
  survey-data.json     — machine-readable, feeds schema design

Usage:
  # From your vista-docs repo root:
  python3 scripts/corpus_survey.py --manifest scripts/manifest.json

  # Or point directly at a directory of docx files:
  python3 scripts/corpus_survey.py --docx-dir raw/cprs

  # Limit to one package:
  python3 scripts/corpus_survey.py --manifest scripts/manifest.json --pkg cprs

  # Verbose: print per-file details as they process
  python3 scripts/corpus_survey.py --manifest scripts/manifest.json --verbose
"""

import json
import re
import sys
import argparse
from pathlib import Path
from collections import Counter, defaultdict
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.  Run:  pip install python-docx")
    sys.exit(1)


# =============================================================================
# 1. CONSTANTS — pattern matchers for VA-specific element types
# =============================================================================

# Styles that are definitively headings in VA documents
HEADING_STYLE_RE = re.compile(r'^Heading\s+\d', re.I)

# TOC paragraph styles
TOC_STYLE_RE = re.compile(r'^TOC\s+\d', re.I)

# Revision history table signals — column header keywords
REVISION_HEADER_KEYWORDS = {"date", "patch", "revision", "description",
                             "change", "pages", "page", "author", "version"}

# Stop words to exclude from header token classification.
# These are so common they add noise without signal.
STOP_WORDS = {
    "a", "an", "the", "and", "or", "of", "in", "to", "for",
    "are", "be", "is", "by", "as", "at", "on", "no", "all",
    "can", "also", "also", "new", "not", "if", "it", "its",
    "was", "with", "from", "that", "this", "then", "than",
    "has", "have", "had", "will", "would", "may", "must",
    "about", "after", "before", "when", "which", "what",
    "only", "any", "each", "both", "do", "does", "used",
    "add", "x", "&", "-", "#", "1", "2", "3", "4", "5",
}

# Table type signals — keywords found in header rows → semantic table type
TABLE_TYPE_SIGNALS = {
    "revision":              {"revision", "patch", "date", "description", "author"},
    "menu-option":           {"menu", "option", "synonym", "lock", "exit"},
    "parameter-list":        {"parameter", "value", "precedence", "entity"},
    "fileman":               {"fileman", "global", "field", "cross-reference"},
    "fileman-file-registry": {"global", "number", "file", "name"},
    "fileman-access":        {"dd", "del", "laygo", "rd", "wr", "access"},
    "m-variable":            {"translation", "variable", "alpha", "call"},
    "security-key":          {"security", "key", "lock"},
    "rpc":                   {"rpc", "tag", "routine", "remote"},
    "pre-post-install":      {"pre-install", "post-install", "routine", "required"},
    "hl7-segment":           {"segment", "hl7", "field", "sequence", "message",
                              "dt", "element", "len"},  # ADT-style HL7 tables
    "api-function":          {"api", "hlo", "entry", "subscript", "pass-by-reference",
                              "pass-by-value", "$$"},
    "keyboard-shortcuts":    {"keystroke", "keyboard", "shortcut", "navigation", "action", "key"},
    "comparison":            {"vs", "versus", "difference", "oe", "rr", "cprs", "version"},
    "checklist":             {"checklist", "readiness", "task", "done", "complete", "action"},
    "acronym-glossary":      {"acronym", "meaning", "abbreviation", "term", "definition"},
    "reminder-definition":   {"reminder", "dialog", "finding", "clinical", "resolution",
                              "frequency", "reminder-definition"},
    "field-mapping":         {"mapping", "url", "source", "target", "interface",
                              "field", "name", "description"},
    "error-code":            {"error", "code", "description", "message", "cause",
                              "resolution", "number"},
}

# Per-type minimum keyword match threshold.
# Most types require 2 matches to avoid false positives on generic headers.
# Types with highly distinctive single tokens (rpc, api-function, fileman-access)
# use threshold=1.  fileman-file-registry needs 3 because its tokens are common words.
MATCH_THRESHOLD = {
    "revision":              2,
    "menu-option":           2,
    "parameter-list":        2,
    "fileman":               2,
    "fileman-file-registry": 3,
    "fileman-access":        2,   # dd+del or laygo+rd are distinctive pairs
    "m-variable":            2,
    "security-key":          2,
    "rpc":                   1,   # "rpc" alone is unambiguous
    "pre-post-install":      2,
    "hl7-segment":           3,   # dt+element+len are common words — require 3
    "api-function":          1,   # "hlo" or "$$" alone is unambiguous
    "keyboard-shortcuts":    2,
    "comparison":            2,
    "checklist":             2,
    "acronym-glossary":      2,
    "reminder-definition":   3,   # reminder+dialog+finding — all distinctive together
    "field-mapping":         3,   # mapping+url+field — avoid firing on generic tables
    "error-code":            3,   # error+code+description — common words, need 3
}


# Callout patterns in paragraph text
CALLOUT_RE = re.compile(
    r'^\s*(?P<type>NOTE|WARNING|CAUTION|IMPORTANT|REMINDER|DISCLAIMER)'
    r'\s*[:–—]',
    re.I
)

# Terminal / VistA session heuristics — short lines, starts with common prompts
TERMINAL_SIGNALS = [
    re.compile(r'^\s*Select\s+\w'),            # VistA menu prompt
    re.compile(r'^\s*>>\s'),                   # MUMPS-style output
    re.compile(r'^\s*D\s+\^[A-Z]'),            # M entrypoint call
    re.compile(r'^\s*\w+\^[A-Z]+\b'),          # M global reference
    re.compile(r'^\s*\[\w+\]\s*$'),            # bracketed option
]

# "See page N" / "See section III" — cross-reference in body text
PAGE_REF_RE = re.compile(
    r'\b(?:see|refer(?:\s+to)?|on)\s+(?:page|section|chapter|p\.)\s*[\dIVXivx]',
    re.I
)
EXTERNAL_REF_RE = re.compile(
    r'\b(?:see|refer(?:\s+to)?)\s+the\s+(?:[A-Za-z][a-zA-Z0-9\s/]+?)'
    r'(?:\s+(?:technical\s+)?(?:manual|guide|handbook|document|reference))',
    re.I
)

# Index section detector — single-letter headings (## A, ## B …)
INDEX_HEAD_RE = re.compile(r'^[A-Z]$')


# =============================================================================
# 2. HELPERS
# =============================================================================

def para_full_text(para) -> str:
    """Full text of a paragraph, joining all runs."""
    return "".join(r.text for r in para.runs).strip()


def para_is_all_bold(para) -> bool:
    """True if every non-empty run in the paragraph is bold."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)


def table_header_tokens(table) -> frozenset:
    """
    Extract text tokens from the first row (and optionally second row if the
    first is a merged caption row) of a table. Returns lowercased frozenset
    with stop words removed.
    """
    if not table.rows:
        return frozenset()
    header_text = " ".join(
        cell.text.strip().lower()
        for cell in table.rows[0].cells
    )
    tokens = frozenset(re.split(r'[\s,/|]+', header_text))
    return tokens - STOP_WORDS


def classify_table(header_tokens: frozenset, rows: int = 0) -> str:
    """
    Return the best matching table type label from TABLE_TYPE_SIGNALS,
    or 'reference-table' if nothing scores at or above its per-type threshold.

    Uses MATCH_THRESHOLD (per-type minimums) rather than a single global
    threshold, so types with highly distinctive single tokens (rpc, api-function)
    can fire on one match while generic types still require two.

    rows: total row count including header — api-function requires >1 data row
    to avoid over-firing on single-row HL7 tables.
    """
    best_type  = "reference-table"
    best_score = 0
    for ttype, keywords in TABLE_TYPE_SIGNALS.items():
        # api-function tables must have at least 2 rows (header + 1 data row)
        if ttype == "api-function" and rows <= 1:
            continue
        score     = len(header_tokens & keywords)
        threshold = MATCH_THRESHOLD.get(ttype, 2)
        if score >= threshold and score > best_score:
            best_score = score
            best_type  = ttype
    return best_type


def is_pseudo_heading(para, body_font_size_pt: float) -> bool:
    """
    Heuristic: paragraph is a visual heading but not styled as Heading N.
    Criteria: all-bold, short (≤ 120 chars), not a table cell, font larger
    than body or explicitly large.
    """
    text = para_full_text(para)
    if not text or len(text) > 120:
        return False
    if not para_is_all_bold(para):
        return False
    # Check if it's inside a table — para.part is the document part, not a cell
    # We detect table paras by checking parent element tag
    parent_tag = para._element.getparent().tag if para._element.getparent() is not None else ""
    if "tc" in parent_tag:      # w:tc = table cell
        return False
    return True


def detect_terminal_block(para) -> bool:
    """Heuristic: paragraph looks like a VistA terminal session line."""
    text = para_full_text(para)
    if not text:
        return False
    style_name = para.style.name if para.style else ""
    # Explicit style signals
    if any(k in style_name.lower() for k in ("code", "mono", "courier", "fixed", "terminal")):
        return True
    # Content signals
    return any(sig.search(text) for sig in TERMINAL_SIGNALS)


def body_font_size(doc) -> float:
    """Estimate body text font size in points from Normal style."""
    try:
        normal = doc.styles["Normal"]
        sz = normal.font.size
        if sz:
            return sz.pt
    except Exception:
        pass
    return 11.0  # Word default


# =============================================================================
# 3. SINGLE-DOCUMENT ANALYSIS
# =============================================================================

def analyse_document(docx_path: Path, verbose: bool = False) -> dict:
    """
    Analyse one .docx file and return a structured survey dict.
    """
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return {"path": str(docx_path), "error": str(e)}

    result = {
        "path":               str(docx_path),
        "filename":           docx_path.name,
        "para_count":         0,
        "style_counts":       Counter(),    # style.name → count
        "heading_styles":     Counter(),    # "Heading 1" etc → count
        "toc_styles":         Counter(),    # "TOC 1" etc → count
        "pseudo_headings":    [],           # list of {text, style}
        "callouts":           Counter(),    # NOTE/WARNING/CAUTION → count
        "terminal_lines":     0,
        "page_refs_internal": 0,
        "page_refs_external": [],           # list of matched doc names
        "has_index":          False,
        "has_toc_field":      False,
        "has_revision_table": False,
        "tables":             [],           # list of {rows, cols, header_tokens, type}
        "table_type_counts":  Counter(),
        "unique_styles":      set(),
    }

    bfs = body_font_size(doc)

    # -----------------------------------------------------------------
    # 3a. Paragraph pass
    # -----------------------------------------------------------------
    for para in doc.paragraphs:
        result["para_count"] += 1
        style_name = para.style.name if para.style else "Default"
        result["style_counts"][style_name] += 1
        result["unique_styles"].add(style_name)

        # Heading styles
        if HEADING_STYLE_RE.match(style_name):
            result["heading_styles"][style_name] += 1
            text = para_full_text(para)
            # Index detector: single-letter heading
            if INDEX_HEAD_RE.match(text.strip()):
                result["has_index"] = True

        # TOC styles
        elif TOC_STYLE_RE.match(style_name):
            result["toc_styles"][style_name] += 1
            result["has_toc_field"] = True

        else:
            text = para_full_text(para)

            # Callouts
            cm = CALLOUT_RE.match(text)
            if cm:
                result["callouts"][cm.group("type").upper()] += 1

            # Terminal sessions
            if detect_terminal_block(para):
                result["terminal_lines"] += 1

            # Cross-references
            if PAGE_REF_RE.search(text):
                result["page_refs_internal"] += 1
            ext = EXTERNAL_REF_RE.search(text)
            if ext:
                result["page_refs_external"].append(ext.group(0)[:80])

            # Pseudo-headings
            if is_pseudo_heading(para, bfs):
                result["pseudo_headings"].append({
                    "text":  text[:100],
                    "style": style_name,
                })

    # -----------------------------------------------------------------
    # 3b. Table pass
    # -----------------------------------------------------------------
    for table in doc.tables:
        rows = len(table.rows)
        cols = max((len(row.cells) for row in table.rows), default=0)
        header_tokens = table_header_tokens(table)

        # Single-column tables are degenerate — likely bulleted lists exported
        # from Word as 1-col tables. Flag separately, skip type classification.
        if cols <= 1:
            ttype = "single-column"
        else:
            ttype = classify_table(header_tokens, rows=rows)

            # Revision table heuristic
            if len(header_tokens & REVISION_HEADER_KEYWORDS) >= 3:
                result["has_revision_table"] = True
                ttype = "revision"

        result["tables"].append({
            "rows":          rows,
            "cols":          cols,
            "header_tokens": sorted(header_tokens)[:12],   # cap for readability
            "type":          ttype,
        })
        result["table_type_counts"][ttype] += 1

    # Convert sets to lists for JSON serialisation
    result["unique_styles"] = sorted(result["unique_styles"])

    if verbose:
        _print_doc_summary(result)

    return result


def _print_doc_summary(r: dict):
    """Print a one-page summary of a document analysis result."""
    print(f"\n  {'─'*60}")
    print(f"  {r['filename']}")
    print(f"  {'─'*60}")
    print(f"  Paragraphs : {r['para_count']}")
    print(f"  Headings   : {sum(r['heading_styles'].values())} "
          f"({', '.join(f'{s}={n}' for s,n in sorted(r['heading_styles'].items()))})")
    print(f"  TOC entries: {sum(r['toc_styles'].values())}")
    print(f"  Tables     : {len(r['tables'])} "
          f"({', '.join(f'{t}={n}' for t,n in r['table_type_counts'].most_common())})")
    print(f"  Pseudo-hdgs: {len(r['pseudo_headings'])}")
    print(f"  Callouts   : {dict(r['callouts'])}")
    print(f"  TermLines  : {r['terminal_lines']}")
    print(f"  PageRefs   : {r['page_refs_internal']} internal, "
          f"{len(r['page_refs_external'])} external")
    print(f"  Flags      : "
          f"{'has_toc ' if r['has_toc_field'] else ''}"
          f"{'has_revtable ' if r['has_revision_table'] else ''}"
          f"{'has_index' if r['has_index'] else ''}")
    print(f"  Top styles : {r['style_counts'].most_common(6)}")


# =============================================================================
# 4. CORPUS AGGREGATION
# =============================================================================

def aggregate(doc_results: list) -> dict:
    """
    Roll up per-document results into corpus-level statistics.
    """
    ok = [r for r in doc_results if "error" not in r]
    errors = [r for r in doc_results if "error" in r]

    corpus_styles       = Counter()
    corpus_table_types  = Counter()
    corpus_callouts     = Counter()
    all_styles_by_doc   = {}   # style → set of filenames where it appears
    pseudo_heading_texts = Counter()
    external_refs       = Counter()
    total_terminal      = 0
    total_page_refs     = 0
    docs_with_toc       = 0
    docs_with_revision  = 0
    docs_with_index     = 0
    docs_with_pseudo    = 0

    for r in ok:
        for style, count in r["style_counts"].items():
            corpus_styles[style] += count
            all_styles_by_doc.setdefault(style, set()).add(r["filename"])

        for ttype, count in r["table_type_counts"].items():
            corpus_table_types[ttype] += count

        for ctype, count in r["callouts"].items():
            corpus_callouts[ctype] += count

        for ph in r["pseudo_headings"]:
            pseudo_heading_texts[ph["text"]] += 1

        for ext in r["page_refs_external"]:
            external_refs[ext] += 1

        total_terminal  += r["terminal_lines"]
        total_page_refs += r["page_refs_internal"]

        if r["has_toc_field"]:      docs_with_toc      += 1
        if r["has_revision_table"]: docs_with_revision += 1
        if r["has_index"]:          docs_with_index    += 1
        if r["pseudo_headings"]:    docs_with_pseudo   += 1

    n = len(ok)

    # Styles that appear in only one document vs. corpus-wide standards
    one_off_styles    = {s for s, docs in all_styles_by_doc.items() if len(docs) == 1}
    standard_styles   = {s for s, docs in all_styles_by_doc.items() if len(docs) == n}
    common_styles     = {s for s, docs in all_styles_by_doc.items()
                         if 1 < len(docs) < n}

    return {
        "documents_ok":         n,
        "documents_error":      len(errors),
        "errors":               [{"path": e["path"], "error": e["error"]} for e in errors],

        # Style inventory
        "corpus_style_counts":  dict(corpus_styles.most_common()),
        "standard_styles":      sorted(standard_styles),   # in EVERY doc
        "common_styles":        sorted(common_styles),     # in >1 but not all docs
        "one_off_styles":       sorted(one_off_styles),    # in exactly 1 doc

        # Table types
        "table_type_counts":    dict(corpus_table_types.most_common()),

        # Callout types
        "callout_counts":       dict(corpus_callouts),

        # Pseudo-headings
        "pseudo_heading_count": sum(len(r["pseudo_headings"]) for r in ok),
        "pseudo_heading_top":   dict(pseudo_heading_texts.most_common(30)),

        # Cross-references
        "total_internal_page_refs":  total_page_refs,
        "total_terminal_lines":      total_terminal,
        "external_doc_refs":         dict(external_refs.most_common(20)),

        # Structural flags
        "docs_with_toc":             docs_with_toc,
        "docs_with_revision_table":  docs_with_revision,
        "docs_with_index":           docs_with_index,
        "docs_with_pseudo_headings": docs_with_pseudo,

        # Per-doc summaries for the report
        "per_doc": [
            {
                "filename":          r["filename"],
                "path":              r["path"],
                "para_count":        r["para_count"],
                "heading_count":     sum(r["heading_styles"].values()),
                "toc_entry_count":   sum(r["toc_styles"].values()),
                "table_count":       len(r["tables"]),
                "table_types":       dict(r["table_type_counts"]),
                "pseudo_headings":   len(r["pseudo_headings"]),
                "callouts":          dict(r["callouts"]),
                "terminal_lines":    r["terminal_lines"],
                "page_refs":         r["page_refs_internal"],
                "has_toc":           r["has_toc_field"],
                "has_revision":      r["has_revision_table"],
                "has_index":         r["has_index"],
                "unique_style_count": len(r["unique_styles"]),
                "tables_detail":     r["tables"],
            }
            for r in ok
        ],
    }


# =============================================================================
# 5. REPORT RENDERING
# =============================================================================

def render_report(corpus: dict, pkg_filter: str | None) -> str:
    lines = []
    w = lines.append

    pkg_label = f"package={pkg_filter}" if pkg_filter else "all packages"
    w("=" * 72)
    w(f"VDL CORPUS SURVEY REPORT — {pkg_label}")
    w(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    w("=" * 72)

    w(f"\nDocuments analysed : {corpus['documents_ok']}")
    w(f"Errors             : {corpus['documents_error']}")
    if corpus["errors"]:
        for e in corpus["errors"]:
            w(f"  ERROR  {e['path']}: {e['error']}")

    # ── Structural flags ──────────────────────────────────────────────
    w("\n" + "─" * 72)
    w("STRUCTURAL FEATURES (count of documents)")
    w("─" * 72)
    n = corpus["documents_ok"]
    w(f"  Has TOC field                : {corpus['docs_with_toc']:3d} / {n}")
    w(f"  Has revision history table   : {corpus['docs_with_revision_table']:3d} / {n}")
    w(f"  Has back-matter index        : {corpus['docs_with_index']:3d} / {n}")
    w(f"  Has pseudo-headings          : {corpus['docs_with_pseudo_headings']:3d} / {n}")

    # ── Table types ──────────────────────────────────────────────────
    w("\n" + "─" * 72)
    w("TABLE TYPES DETECTED")
    w("─" * 72)
    for ttype, count in sorted(corpus["table_type_counts"].items(),
                                key=lambda x: -x[1]):
        w(f"  {ttype:<25} {count:4d}")

    # ── Callout types ────────────────────────────────────────────────
    w("\n" + "─" * 72)
    w("CALLOUT TYPES (NOTE / WARNING / CAUTION etc.)")
    w("─" * 72)
    for ctype, count in sorted(corpus["callout_counts"].items(),
                                key=lambda x: -x[1]):
        w(f"  {ctype:<20} {count:4d}")

    # ── Cross-references ─────────────────────────────────────────────
    w("\n" + "─" * 72)
    w("CROSS-REFERENCES")
    w("─" * 72)
    w(f"  Internal page refs           : {corpus['total_internal_page_refs']}")
    w(f"  Terminal/VistA session lines : {corpus['total_terminal_lines']}")
    w(f"  External doc refs (top 10):")
    for ref, count in list(corpus["external_doc_refs"].items())[:10]:
        w(f"    {count:3d}×  {ref}")

    # ── Pseudo-headings ──────────────────────────────────────────────
    w("\n" + "─" * 72)
    w(f"PSEUDO-HEADINGS — total {corpus['pseudo_heading_count']} "
      f"across {corpus['docs_with_pseudo_headings']} docs")
    w("Top 20 recurring pseudo-heading texts:")
    w("─" * 72)
    for text, count in list(corpus["pseudo_heading_top"].items())[:20]:
        w(f"  {count:3d}×  {text}")

    # ── Paragraph styles ─────────────────────────────────────────────
    w("\n" + "─" * 72)
    w("PARAGRAPH STYLES — corpus frequency")
    w("─" * 72)
    for style, count in list(corpus["corpus_style_counts"].items())[:40]:
        w(f"  {count:6d}  {style}")

    # ── Style universality ───────────────────────────────────────────
    w("\n" + "─" * 72)
    w(f"STANDARD STYLES (appear in ALL {n} documents) — these form the base schema")
    w("─" * 72)
    for s in corpus["standard_styles"]:
        w(f"  {s}")

    w("\n" + "─" * 72)
    w("ONE-OFF STYLES (appear in exactly 1 document) — candidates for schema 'other'")
    w("─" * 72)
    for s in corpus["one_off_styles"][:30]:
        w(f"  {s}")
    if len(corpus["one_off_styles"]) > 30:
        w(f"  ... and {len(corpus['one_off_styles']) - 30} more")

    # ── Per-document table ───────────────────────────────────────────
    w("\n" + "─" * 72)
    w("PER-DOCUMENT SUMMARY")
    w("─" * 72)
    w(f"  {'Document':<45} {'§':>4} {'Tbl':>4} {'Ph':>4} {'PgRef':>6} {'Flags'}")
    w(f"  {'-'*45} {'─'*4} {'─'*4} {'─'*4} {'─'*6} {'─'*12}")
    for d in corpus["per_doc"]:
        flags = ""
        if d["has_toc"]:      flags += "T"
        if d["has_revision"]: flags += "R"
        if d["has_index"]:    flags += "I"
        w(f"  {d['filename']:<45} "
          f"{d['heading_count']:>4} "
          f"{d['table_count']:>4} "
          f"{d['pseudo_headings']:>4} "
          f"{d['page_refs']:>6} "
          f"  {flags}")
    w("  Flags: T=has TOC field  R=has revision table  I=has back-matter index")

    w("\n" + "=" * 72)
    w("END OF REPORT")
    w("=" * 72)

    return "\n".join(lines)


# =============================================================================
# 6. FILE DISCOVERY
# =============================================================================

def find_docx_files(manifest_path: Path, pkg_filter: str | None,
                    repo_root: Path) -> list[Path]:
    """
    Find all fetched base-document .docx files from the manifest.
    Falls back to walking raw/ if local_path is not populated.
    """
    with open(manifest_path, encoding="utf-8") as f:
        manifest = json.load(f)

    paths = []
    for pkg_key, pkg in manifest["packages"].items():
        if pkg_filter and pkg_key != pkg_filter:
            continue
        for doc in pkg["documents"]:
            if not doc.get("is_base"):
                continue
            # Try local_path first (set by fetch.py)
            lp = doc.get("local_path")
            if lp:
                p = repo_root / lp
                if p.exists() and p.suffix.lower() == ".docx":
                    paths.append(p)
                    continue
            # Fallback: guess from raw/ directory layout
            # Pattern: raw/{pkg_key}/{docx_filename}
            fname = doc.get("docx_filename") or doc.get("pdf_filename", "")
            if fname:
                # docx_filename may be the full original name
                guessed = repo_root / "raw" / pkg_key / fname
                if guessed.exists():
                    paths.append(guessed)

    return paths


def find_docx_in_dir(docx_dir: Path) -> list[Path]:
    """Walk a directory and return all .docx files found recursively."""
    return sorted(docx_dir.rglob("*.docx"))


# =============================================================================
# 7. MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Corpus survey: element-type frequency analysis across VDL DOCX files."
    )
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("--manifest", help="Path to manifest.json")
    src.add_argument("--docx-dir", help="Directory containing .docx files (scanned recursively)")
    parser.add_argument("--pkg",     default=None,
                        help="Limit to one package key, e.g. cprs")
    parser.add_argument("--verbose", action="store_true",
                        help="Print per-file summary as each file is processed")
    parser.add_argument("--out-dir", default=None,
                        help="Directory for output files (default: same dir as manifest, "
                             "or current dir for --docx-dir)")
    args = parser.parse_args()

    # ── Locate docx files ────────────────────────────────────────────
    if args.manifest:
        manifest_path = Path(args.manifest).resolve()
        repo_root     = manifest_path.parent.parent
        docx_files    = find_docx_files(manifest_path, args.pkg, repo_root)
        out_dir       = Path(args.out_dir) if args.out_dir else manifest_path.parent
    else:
        docx_dir   = Path(args.docx_dir).resolve()
        docx_files = find_docx_in_dir(docx_dir)
        out_dir    = Path(args.out_dir) if args.out_dir else Path.cwd()

    if not docx_files:
        print(
            "No .docx files found.\n\n"
            "If using --manifest: run fetch.py first so local_path fields are populated,\n"
            "or place .docx files in raw/{pkg}/ and re-run.\n\n"
            "If using --docx-dir: check the path is correct and contains .docx files."
        )
        sys.exit(1)

    print(f"Found {len(docx_files)} .docx file(s) to analyse.")

    # ── Analyse each file ────────────────────────────────────────────
    doc_results = []
    for i, path in enumerate(docx_files, 1):
        print(f"  [{i:2d}/{len(docx_files)}] {path.name}")
        result = analyse_document(path, verbose=args.verbose)
        doc_results.append(result)

    # ── Aggregate ────────────────────────────────────────────────────
    corpus = aggregate(doc_results)

    # ── Outputs ──────────────────────────────────────────────────────
    out_dir.mkdir(parents=True, exist_ok=True)
    pkg_suffix = f"-{args.pkg}" if args.pkg else ""

    report_path = out_dir / f"survey-report{pkg_suffix}.txt"
    data_path   = out_dir / f"survey-data{pkg_suffix}.json"

    report_text = render_report(corpus, args.pkg)
    report_path.write_text(report_text, encoding="utf-8")

    # Make style counter JSON-serialisable
    corpus_json = {
        **corpus,
        "per_doc": [
            {**d, "table_types": dict(d["table_types"])}
            for d in corpus["per_doc"]
        ]
    }
    data_path.write_text(
        json.dumps(corpus_json, indent=2, default=str),
        encoding="utf-8"
    )

    print(f"\nDone.")
    print(f"  Report : {report_path}")
    print(f"  Data   : {data_path}")
    print()

    # ── Quick console summary ────────────────────────────────────────
    print("── Quick summary ──────────────────────────────────────────────")
    print(f"  Docs analysed          : {corpus['documents_ok']}")
    print(f"  Standard styles        : {len(corpus['standard_styles'])}")
    print(f"  One-off styles         : {len(corpus['one_off_styles'])}")
    print(f"  Total tables           : {sum(corpus['table_type_counts'].values())}")
    print(f"  Pseudo-headings        : {corpus['pseudo_heading_count']}")
    print(f"  Docs with TOC field    : {corpus['docs_with_toc']}")
    print(f"  Docs with rev. table   : {corpus['docs_with_revision_table']}")
    print(f"  Internal page refs     : {corpus['total_internal_page_refs']}")
    print(f"  External doc refs      : {sum(corpus['external_doc_refs'].values())}")
    print()


if __name__ == "__main__":
    main()