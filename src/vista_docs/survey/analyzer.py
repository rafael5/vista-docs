"""I/O thin layer: walk corpus DOCX files and run detectors over them."""
from __future__ import annotations

import json
import logging
from pathlib import Path

from vista_docs.survey.detectors import classify_style, classify_table, detect_callout

logger = logging.getLogger(__name__)


def survey_package(pkg_dir: Path) -> dict:
    """Survey all DOCX files in a package directory. Returns analysis dict."""
    try:
        import docx
    except ImportError as exc:
        raise ImportError("python-docx is required for survey: pip install python-docx") from exc

    results: dict = {"package": pkg_dir.name, "files": [], "styles": {}, "tables": {}, "callouts": {}}

    for docx_path in sorted(pkg_dir.glob("*.docx")):
        logger.info("Surveying %s", docx_path.name)
        try:
            doc = docx.Document(str(docx_path))
            _survey_document(doc, docx_path.name, results)
        except Exception as exc:
            logger.warning("Failed to survey %s: %s", docx_path.name, exc)

    return results


def _survey_document(doc, filename: str, results: dict) -> None:
    """Accumulate survey data from a single python-docx Document."""
    file_entry: dict = {"name": filename, "paragraphs": 0, "tables": 0, "callouts": 0}

    for para in doc.paragraphs:
        style = para.style.name if para.style else "Normal"
        results["styles"][style] = results["styles"].get(style, 0) + 1
        file_entry["paragraphs"] += 1
        if detect_callout(para.text):
            file_entry["callouts"] += 1
            prefix = detect_callout(para.text)
            results["callouts"][prefix] = results["callouts"].get(prefix, 0) + 1

    for table in doc.tables:
        file_entry["tables"] += 1
        headers = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_type = classify_table(headers)
        results["tables"][table_type] = results["tables"].get(table_type, 0) + 1

    results["files"].append(file_entry)
